from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Relatorio_Achados_Revisao_Codigo_2026-04-29.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Relatorio de Revisao Tecnica")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(97, 0, 5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Projeto CarneUp / Junior Prime Beef")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(90, 64, 60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Data: {date(2026, 4, 29).strftime('%d/%m/%Y')}")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_paragraph("")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(6.6)
    cell = box.cell(0, 0)
    cell.width = Inches(6.6)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, "F9E7E7")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Objetivo: consolidar os principais erros identificados na revisao do codigo, "
        "com foco em banco de dados, integracao front/back, logica funcional e riscos "
        "para a operacao e para a pipeline."
    )
    r.font.size = Pt(11)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    r = p.add_run(text)
    if level == 1:
        r.font.color.rgb = RGBColor(97, 0, 5)
    else:
        r.font.color.rgb = RGBColor(127, 29, 29)
    return p


def add_summary_table(doc):
    add_heading(doc, "Resumo Executivo", level=1)
    p = doc.add_paragraph(
        "A revisao encontrou quatro problemas principais. Dois deles sao de severidade "
        "alta e afetam diretamente a confiabilidade do ambiente e a aderencia do sistema "
        "ao fluxo real de uso."
    )
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["ID", "Severidade", "Area", "Resumo"]
    for cell, text in zip(hdr, headers):
        set_cell_shading(cell, "610005")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows = [
        ("F1", "P1", "Banco / Flyway", "Cadeia de migrations nao sobe banco novo do zero."),
        ("F2", "P1", "Frontend / Integracao", "Telas principais ainda usam mock e nao persistem no backend."),
        ("F3", "P2", "Logica de Compra", "Edicao usa id do item local como se fosse id do produto."),
        ("F4", "P2", "Teste / Onboarding", "Teste do gerador de usuario ainda sugere nivel de acesso invalido."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            r.font.size = Pt(10)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_finding(doc, fid, title, severity, file_ref, impact, evidence, recommendation):
    add_heading(doc, f"{fid} - {title}", level=2)

    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    labels = [("Severidade", severity), ("Arquivo", file_ref)]
    for idx, (label, value) in enumerate(labels):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        set_cell_shading(left, "F3F3F3")
        left.text = label
        right.text = value
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        left.paragraphs[0].runs[0].bold = True
        for cell in (left, right):
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run("Impacto: ")
    r.bold = True
    p.add_run(impact)

    p = doc.add_paragraph()
    r = p.add_run("Evidencia: ")
    r.bold = True
    p.add_run(evidence)

    p = doc.add_paragraph()
    r = p.add_run("Recomendacao: ")
    r.bold = True
    p.add_run(recommendation)


def add_conclusion(doc):
    add_heading(doc, "Conclusao", level=1)
    p = doc.add_paragraph(
        "O projeto ja apresenta base funcional de autenticacao, testes de controller e "
        "pipeline inicial, mas ainda ha risco relevante para deploy limpo e para uso real "
        "do produto. O proximo ciclo de trabalho deve priorizar a consolidacao das "
        "migrations do banco e a integracao real das telas de estoque, compras e dashboard "
        "com os endpoints do backend."
    )
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.add_run("Prioridade sugerida para o time:").bold = True
    for item in [
        "Corrigir a estrategia de migrations para bootstrap de banco novo.",
        "Integrar StockView, PurchaseView e DashboardView com o backend real.",
        "Corrigir o fluxo de edicao de compras no frontend.",
        "Atualizar o teste do gerador de usuario inicial para refletir ADM.",
    ]:
        doc.add_paragraph(item, style="List Bullet")


doc = Document()
for section in doc.sections:
    set_page_margins(section)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(15)
styles["Heading 2"].font.name = "Aptos Display"
styles["Heading 2"].font.size = Pt(12)

add_title(doc)
doc.add_page_break()
add_summary_table(doc)
add_finding(
    doc,
    "F1",
    "Flyway nao consegue criar banco novo com a cadeia atual de migrations",
    "P1",
    "Source/Server/SpringBootApp/src/main/resources/migrations/V2_DATABASE-MODEL.sql:3-48",
    "Impede bootstrap limpo do ambiente e torna o deploy dependente de um banco legado ja existente.",
    "A migration V2 volta a executar CREATE TABLE para entidades que ja existem em V1, como Produto, Categoria, Marca, Usuario, Compra e Venda.",
    "Reorganizar as migrations para que a sequencia completa consiga subir uma base vazia, ou consolidar um baseline novo e apos isso manter somente migrations incrementais.",
)
add_finding(
    doc,
    "F2",
    "Frontend principal ainda nao usa dados reais do backend",
    "P1",
    "Source/Client/carneup-frontend/src/views/StockView.jsx:133-190",
    "Mesmo com login funcional, o usuario nao opera sobre o Postgres nem sobre os endpoints principais do sistema.",
    "StockView, PurchaseView e DashboardView ainda usam mocks, console.log e alert em vez de chamadas HTTP para /products, /purchases e /sales.",
    "Criar camada de servicos para estoque, compras e vendas, substituir mocks locais por respostas reais da API e tratar loading, erro e persistencia.",
)
add_finding(
    doc,
    "F3",
    "Fluxo de edicao de compra usa id incorreto",
    "P2",
    "Source/Client/carneup-frontend/src/views/PurchaseView.jsx:41-63",
    "A edicao do item no carrinho nao reconstrói corretamente o formulario, comprometendo a manutencao da compra antes da gravacao.",
    "O item local do carrinho recebe id baseado em Date.now(), mas esse valor depois e reutilizado como selectedProductId.",
    "Separar claramente itemId e productId dentro do carrinho, e reconstruir o estado do formulario com base no productId real do produto selecionado.",
)
add_finding(
    doc,
    "F4",
    "Teste do gerador de usuario inicial esta desatualizado",
    "P2",
    "Source/Server/SpringBootApp/src/test/java/com/example/SpringBootApp/utils/PasswordHashGeneratorTest.java:23-29",
    "Pode induzir o time a criar manualmente usuario com nivel de acesso invalido para o schema atual.",
    "O utilitario principal imprime ADM, mas o teste auxiliar ainda imprime ADMIN.",
    "Atualizar o teste para refletir ADM e evitar ambiguidade no processo de onboarding tecnico.",
)

add_conclusion(doc)
doc.save(OUTPUT)
print(OUTPUT)
